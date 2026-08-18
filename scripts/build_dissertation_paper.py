"""Build an evidence-frozen MSc dissertation DOCX and publication figures.

The caller supplies the Markdown scholarly record because the public repository
does not redistribute the dissertation manuscript. This script applies the
recorded A4, two-column paper conventions and embeds compact figures generated
only from frozen summary tables or explicitly registered values.
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

import matplotlib
import numpy as np

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from matplotlib.lines import Line2D
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch, Patch
from matplotlib.ticker import PercentFormatter

ROOT = Path(__file__).resolve().parents[1]
SRC_DIR = ROOT / "src"
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

from market_dynamics.reporting.dissertation_figures import (  # noqa: E402
    CHANCE_AUC,
    FAMILY_ORDER,
    METHODOLOGY_FIGURE_DATA,
    MODEL_WINDOW_EXCLUSIONS,
    PREVALENCE_EQUALITY_LINE,
    load_appendix_cross_model_results,
    load_asset_prevalence,
    load_core_results,
    load_family_prevalence,
    load_identity_order_results,
    load_simulation_figure_data,
)

DEFAULT_OUTPUT = ROOT / "results" / "dissertation_build" / "verification" / "dissertation.docx"
DEFAULT_ARTIFACTS = ROOT / "results" / "dissertation_build" / "v4"
SIMULATION_TABLE = ROOT / "reports" / "tables" / "prp1_study_a_independent_simulation_results.csv"
IDENTITY_DECOMPOSITION_TABLE = ROOT / "reports" / "tables" / "ifddrp_identity_dynamic_information_decomposition.csv"
CROSS_MODEL_RESULTS_TABLE = ROOT / "reports" / "tables" / "prp1_fixed_cross_model_results.csv"
IDENTITY_SWAP_TABLE = ROOT / "reports" / "tables" / "phase6_identity_swap_results.csv"
TEMPORAL_ORDER_TABLE = ROOT / "reports" / "tables" / "phase6_temporal_order_destruction.csv"
WINDOW_ENDPOINT_PREVALENCE_TABLE = (
    ROOT
    / "src"
    / "market_dynamics"
    / "reporting"
    / "data"
    / "final_model_window_prevalence.csv"
)

matplotlib.rcParams["svg.hashsalt"] = "financial-dynamics-dissertation"
matplotlib.rcParams["svg.fonttype"] = "none"

TITLE = "Interpretable Transformer Models for Financial Time Series Forecasting: Discovering Emergent Market Dynamics"
AUTHOR = "Husaam Ateeq"

INK = "#17252A"
TEAL = "#2A7F8E"
RED = "#B94A48"
GOLD = "#C38D2E"
BLUE = "#486A9A"
LIGHT = "#E9F0F1"
MID = "#70838A"

FAMILY_COLOURS = {
    "Equities": TEAL,
    "Bonds": BLUE,
    "Commodities": GOLD,
    "FX": "#7568A6",
    "Crypto": RED,
    "Real assets": "#5C7D62",
}
FAMILY_LABELS = {
    "Equities": "Equities",
    "Bonds": "Bonds",
    "Commodities": "Commodities",
    "FX": "FX",
    "Crypto": "Crypto",
    "Real assets": "Real-asset\nproxies",
}

FIGURE_CAPTIONS = {
    "methodology": (
        "Figure 1. Forecast construction and corrected chronological evaluation. "
        "Every window ends at the forecast origin; its label uses only the following ten observed sessions."
    ),
    "core_results": (
        "Figure 2. Pooled and pair-weighted within-asset ROC-AUC. The static asset prior is constant "
        "through time within each asset, so its pooled advantage is cross-sectional rather than temporal."
    ),
    "simulation": (
        "Figure 3. Controlled simulation. Static prior heterogeneity raises pooled discrimination without "
        "within-asset skill (top); planted ordered signal raises within-asset AUC and perturbation sensitivity (bottom)."
    ),
}

FIGURE_CAPTIONS_V2 = {
    "methodology": (
        "Figure 1. Forecast design and chronological evaluation. Each 60-session window ends at close t; "
        "only the next ten observed sessions define its label."
    ),
    "simulation": (
        "Figure 3. Controlled simulation. Prior heterogeneity inflates pooled AUC without within-asset skill "
        "(top); planted ordered signal restores within-asset skill and order sensitivity (bottom)."
    ),
}

FIGURE_CAPTIONS_V3 = {
    "methodology": (
        "Figure 1. Forecast design and chronological evaluation. The model receives 60 observed sessions "
        "through close t; only sessions t+1 to t+10 define the adverse-event label."
    ),
    "core_results": (
        "Figure 2. Pooled versus pair-weighted within-asset ROC-AUC. The static asset prior scores 0.824 "
        "pooled but 0.500 within asset because its score never changes through time."
    ),
    "simulation": (
        "Figure 3. Controlled mechanism checks. Static prior heterogeneity inflates pooled AUC without timing "
        "skill (A); planted ordered signal raises within-asset AUC and makes temporal destruction costly (B)."
    ),
}

TABLE_CAPTIONS = {
    1: "Table I. Direct comparison with the closest methodological work.",
    2: "Table II. Direct numerical channels supplied at each of the 60 timesteps.",
    3: "Table III. Cross-model ranking on the corrected held-out test split.",
}

TABLE_CAPTIONS_V3 = {
    **TABLE_CAPTIONS,
    3: "Table III. Cross-model ranking on the corrected historical test split.",
    4: "Table IV. Bounded recovery tests on the corrected historical test split.",
}

EQUATION_RENDER = {
    "X_{i,t} = [x_{i,t-59}, ..., x_{i,t}], where L = 60.": "X(i,t) = [x(i,t-59), ..., x(i,t)], where L = 60.",
    "R_{i,t}^{(10)} = P_{i,t+10}/P_{i,t} - 1,": "R^(10)(i,t) = P(i,t+10) / P(i,t) - 1,",
    "D_{i,t}^{(10)} = min_{1 <= k <= 10}(P_{i,t+k}/P_{i,t} - 1),": "D^(10)(i,t) = min[1 <= k <= 10] {P(i,t+k) / P(i,t) - 1},",
    "V_{i,t}^{future} = sqrt(sum_{k=1}^{10} r_{i,t+k}^2), and V_{i,t}^{hist} = sqrt(10) sd(r_{i,t-19:t}).": "V^future(i,t) = sqrt[sum(k=1..10) r(i,t+k)^2], and V^hist(i,t) = sqrt(10) sd[r(i,t-19:t)].",
    "y_{i,t} = 1{R_{i,t}^{(10)} <= -0.05 OR D_{i,t}^{(10)} <= -0.07 OR V_{i,t}^{future} >= 2 V_{i,t}^{hist}}.": "y(i,t) = 1{R^(10)(i,t) <= -0.05 OR D^(10)(i,t) <= -0.07 OR V^future(i,t) >= 2 V^hist(i,t)}.",
    "AUC_within = sum_i(n_i^+ n_i^- AUC_i) / sum_i(n_i^+ n_i^-).": "AUC_within = sum_i[n_i(+) n_i(-) AUC_i] / sum_i[n_i(+) n_i(-)].",
}

EQUATION_RENDER_V2 = {
    **EQUATION_RENDER,
    "AUC_within = sum_i(n_i^+ n_i^- AUC_i) / sum_i(n_i^+ n_i^-).": "AUC(within) = sum over i {n(i,+) n(i,-) AUC(i)} / sum over i {n(i,+) n(i,-)}.",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--source",
        type=Path,
        required=True,
        help="Path to a locally supplied dissertation Markdown manuscript.",
    )
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--artifacts", type=Path, default=DEFAULT_ARTIFACTS)
    parser.add_argument("--edition", choices=("v1", "v2", "v3", "v4"), help="Figure/layout edition; inferred from source when omitted")
    return parser.parse_args()


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 50, start: int = 65, bottom: int = 50, end: int = 65) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single" if edge in {"top", "bottom", "insideH"} else "nil")
        element.set(qn("w:sz"), "5")
        element.set(qn("w:color"), "809096")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _set_keep_together(paragraph) -> None:
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.widow_control = True


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)


def _configure_section(section, columns: int) -> None:
    section.page_width = Inches(8.268)
    section.page_height = Inches(11.693)
    section.top_margin = Inches(0.375)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.18)
    section.footer_distance = Inches(0.35)

    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(columns))
    cols.set(qn("w:space"), "300")
    cols.set(qn("w:equalWidth"), "1")


def _configure_document(doc: Document) -> None:
    _configure_section(doc.sections[0], columns=1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10.1)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.04
    normal.paragraph_format.space_after = Pt(2.8)
    normal.paragraph_format.widow_control = True

    for name, size, bold, before, after in (
        ("Heading 1", 11.0, True, 6.0, 2.4),
        ("Heading 2", 10.0, True, 4.0, 1.5),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string("17252A")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    if "Equation Compact" not in styles:
        equation = styles.add_style("Equation Compact", WD_STYLE_TYPE.PARAGRAPH)
    else:
        equation = styles["Equation Compact"]
    equation.font.name = "Cambria Math"
    equation._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    equation.font.size = Pt(9.2)
    equation.font.italic = True
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.space_before = Pt(1.2)
    equation.paragraph_format.space_after = Pt(1.8)
    equation.paragraph_format.keep_together = True

    if "Caption Compact" not in styles:
        caption = styles.add_style("Caption Compact", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Caption Compact"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(8.0)
    caption.font.italic = True
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    caption.paragraph_format.space_before = Pt(1.0)
    caption.paragraph_format.space_after = Pt(3.0)
    caption.paragraph_format.keep_together = True

    settings = doc.settings._element
    if settings.find(qn("w:autoHyphenation")) is None:
        hyphenation = OxmlElement("w:autoHyphenation")
        hyphenation.set(qn("w:val"), "true")
        settings.append(hyphenation)

    doc.core_properties.title = TITLE
    doc.core_properties.author = AUTHOR
    doc.core_properties.subject = "MSc Data Science dissertation"
    doc.core_properties.keywords = "financial time series, Transformer, shortcut learning, within-asset evaluation"


def _add_inline_markdown(paragraph, text: str, *, size: float | None = None) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            if size:
                run.font.size = Pt(size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
        if size:
            run.font.size = Pt(size)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        if size:
            run.font.size = Pt(size)


def _make_methodology_figure(path: Path, edition: str) -> None:
    full_budget = edition in {"v3", "v4"}
    data = METHODOLOGY_FIGURE_DATA
    fig_size = (3.38, 4.25) if full_budget else (3.35, 4.15)
    fig, ax = plt.subplots(figsize=fig_size, dpi=320)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    boxes = (
        (0.915, 0.105, f"{data.lookback} observed sessions\nending at close t", TEAL),
        (
            0.755,
            0.135,
            f"{data.numerical_channels} scaled numerical channels\n"
            f"{data.market_channels} market + {data.context_channels} context\n"
            f"+ {data.asset_embedding_dim}-D learned asset embedding",
            BLUE,
        ),
        (
            0.575,
            0.105,
            f"{data.encoder_layers}-layer Transformer\n{data.parameters:,} parameters",
            INK,
        ),
        (
            0.405,
            0.105,
            "Adverse-event probability\n"
            f"for sessions t+1 to t+{data.forecast_horizon}",
            RED,
        ),
    )
    for y, height, label, colour in boxes:
        patch = FancyBboxPatch(
            (0.15, y - height / 2),
            0.70,
            height,
            boxstyle="round,pad=0.008,rounding_size=0.012",
            facecolor=colour,
            edgecolor="white",
            linewidth=0.8,
        )
        ax.add_patch(patch)
        ax.text(
            0.50,
            y,
            label,
            ha="center",
            va="center",
            color="white",
            fontsize=7.7 if full_budget else 7.4,
            weight="bold",
            linespacing=1.18,
        )
    for y1, y2 in ((0.855, 0.829), (0.682, 0.642), (0.518, 0.468)):
        ax.add_patch(
            FancyArrowPatch(
                (0.5, y1),
                (0.5, y2),
                arrowstyle="-|>",
                mutation_scale=8,
                color=MID,
                linewidth=0.9,
            )
        )

    ax.text(
        0.03,
        0.322,
        "Corrected chronological evaluation",
        fontsize=8.1 if full_budget else 7.8,
        weight="bold",
        color=INK,
    )

    timeline_y = 0.247
    timeline_height = 0.060
    first_boundary = 0.535
    second_boundary = 0.775
    purge_width = 0.027
    embargo_width = 0.018
    segments = (
        (0.035, first_boundary - purge_width, TEAL, "Train\n2010–2023"),
        (first_boundary - purge_width, first_boundary, GOLD, ""),
        (first_boundary, first_boundary + embargo_width, "#E3BD69", ""),
        (first_boundary + embargo_width, second_boundary - purge_width, BLUE, "Validation\n2023–2025"),
        (second_boundary - purge_width, second_boundary, GOLD, ""),
        (second_boundary, second_boundary + embargo_width, "#E3BD69", ""),
        (second_boundary + embargo_width, 0.965, RED, "Test\n2025–2026"),
    )
    for left, right, colour, label in segments:
        ax.add_patch(
            FancyBboxPatch(
                (left, timeline_y),
                right - left,
                timeline_height,
                boxstyle="square,pad=0",
                facecolor=colour,
                edgecolor="white",
                linewidth=0.45,
            )
        )
        if label:
            ax.text(
                (left + right) / 2,
                timeline_y + timeline_height / 2,
                label,
                ha="center",
                va="center",
                color="white",
                fontsize=6.6 if full_budget else 6.3,
                weight="bold",
                linespacing=1.05,
            )

    callouts = (
        (0.04, 0.46, "Train → Validation", first_boundary),
        (0.54, 0.96, "Validation → Test", second_boundary),
    )
    for left, right, heading, boundary in callouts:
        ax.plot(
            [boundary, (left + right) / 2],
            [timeline_y, 0.183],
            color=GOLD,
            linewidth=0.8,
            clip_on=False,
        )
        ax.axvline(
            boundary,
            ymin=timeline_y,
            ymax=timeline_y + timeline_height,
            color=INK,
            linewidth=0.75,
        )
        ax.add_patch(
            FancyBboxPatch(
                (left, 0.077),
                right - left,
                0.106,
                boxstyle="round,pad=0.005,rounding_size=0.008",
                facecolor="#F8F3E7",
                edgecolor=GOLD,
                linewidth=0.7,
            )
        )
        ax.text(
            (left + right) / 2,
            0.158,
            heading,
            ha="center",
            va="center",
            fontsize=6.5,
            weight="bold",
            color=INK,
        )
        ax.text(
            (left + right) / 2,
            0.126,
            f"{data.purge_dates}-date purge before",
            ha="center",
            va="center",
            fontsize=6.3,
            color="#8A5A00",
        )
        ax.text(
            (left + right) / 2,
            0.096,
            f"{data.embargo_dates}-date embargo after",
            ha="center",
            va="center",
            fontsize=6.3,
            color="#8A5A00",
        )

    ax.text(
        0.50,
        0.024,
        f"Interval audit: {data.audited_boundary_crossings} split-boundary crossings",
        ha="center",
        fontsize=6.6,
        color=INK,
    )

    fig.subplots_adjust(left=0.02, right=0.98, bottom=0.01, top=0.99)
    _save_publication_figure(fig, path, "Forecast design and chronological evaluation")
    plt.close(fig)


def _make_core_results_figure(path: Path, edition: str) -> None:
    full_budget = edition in {"v3", "v4"}
    results = load_core_results(IDENTITY_DECOMPOSITION_TABLE, CROSS_MODEL_RESULTS_TABLE)
    labels = [result.label for result in results]
    pooled = [result.pooled_auc for result in results]
    within = [result.within_asset_auc for result in results]

    fig_size = (3.38, 2.80) if full_budget else (3.35, 2.70)
    fig, ax = plt.subplots(figsize=fig_size, dpi=320)
    y = list(range(len(labels)))
    offset = 0.18
    pooled_bars = ax.barh(
        [value + offset for value in y],
        pooled,
        height=0.30,
        label="Pooled",
        color=TEAL,
    )
    within_bars = ax.barh(
        [value - offset for value in y],
        within,
        height=0.30,
        label="Within-asset",
        color=GOLD,
    )
    ax.axvline(0.5, color=INK, linewidth=0.8, linestyle="--", zorder=0)
    ax.text(0.503, 3.50, "chance", fontsize=6.3, color=INK, va="bottom")
    ax.set_xlim(0.44, 0.875)
    ax.set_ylim(3.60, -0.55)
    ax.set_xticks([0.5, 0.6, 0.7, 0.8])
    axis_font = 8.0 if full_budget else 7.6
    label_font = 7.8 if full_budget else 7.4
    ax.set_xlabel("ROC-AUC", fontsize=axis_font)
    ax.set_yticks(y, labels, fontsize=label_font)
    ax.tick_params(axis="x", labelsize=7.3 if full_budget else 7.0)
    ax.grid(axis="x", alpha=0.2, linewidth=0.5)
    ax.spines[["top", "right", "left"]].set_visible(False)
    fig.legend(
        (pooled_bars, within_bars),
        ("Pooled", "Within-asset"),
        frameon=False,
        fontsize=7.4 if full_budget else 7.0,
        loc="upper center",
        bbox_to_anchor=(0.60, 0.985),
        ncol=2,
        columnspacing=1.5,
        handlelength=1.6,
    )
    if full_budget:
        ax.get_yticklabels()[0].set_fontweight("bold")
    for row, (pooled_value, within_value) in enumerate(zip(pooled, within, strict=True)):
        value_font = 6.8 if full_budget else 6.4
        ax.text(
            pooled_value + 0.006,
            row + offset,
            f"{pooled_value:.3f}",
            va="center",
            fontsize=value_font,
            color=INK,
        )
        ax.text(
            max(within_value + 0.006, 0.507),
            row - offset,
            f"{within_value:.3f}",
            va="center",
            fontsize=value_font,
            color=INK,
            bbox={"facecolor": "white", "edgecolor": "none", "pad": 0.2, "alpha": 0.88},
        )
    fig.subplots_adjust(left=0.34, right=0.97, bottom=0.14, top=0.85)
    _save_publication_figure(fig, path, "Pooled versus within-asset ROC-AUC")
    plt.close(fig)


def _make_simulation_figure(path: Path, edition: str) -> None:
    full_budget = edition in {"v3", "v4"}
    data = load_simulation_figure_data(SIMULATION_TABLE)

    fig_size = (3.38, 4.75) if full_budget else (3.35, 4.45)
    fig, axes = plt.subplots(2, 1, figsize=fig_size, dpi=320)
    top, bottom = axes
    top.plot(
        data.prior_heterogeneity,
        data.static_prior_pooled_auc,
        marker="o",
        color=RED,
        linewidth=1.8,
        label="Static-prior pooled AUC",
    )
    top.plot(
        data.prior_heterogeneity,
        data.no_signal_within_asset_auc,
        marker="s",
        color=TEAL,
        linewidth=1.8,
        label="Classifier within-asset AUC",
    )
    top.axhline(0.5, color=INK, linewidth=0.7, linestyle="--")
    top.set_ylim(0.45, 0.94)
    axis_font = 8.2 if full_budget else 7.9
    title_font = 8.8 if full_budget else 8.4
    legend_font = 6.9 if full_budget else 6.6
    top.set_ylabel("ROC-AUC", fontsize=axis_font)
    top.set_xlabel("Prior heterogeneity", fontsize=axis_font)
    top.set_title("A. No planted temporal signal", fontsize=title_font, loc="left", weight="bold")
    top.legend(frameon=False, fontsize=legend_font, loc="upper left")

    bottom.plot(
        data.dynamic_signal,
        data.dynamic_within_asset_auc,
        marker="o",
        color=TEAL,
        linewidth=1.8,
        label="Within-asset AUC",
    )
    bottom.plot(
        data.dynamic_signal,
        data.reversal_auc_loss,
        marker="^",
        color=GOLD,
        linewidth=1.8,
        label="AUC loss after reversal",
    )
    bottom.axhline(0.5, color=INK, linewidth=0.7, linestyle="--")
    bottom.set_ylim(-0.03, 0.84)
    bottom.text(
        0.08,
        0.478,
        "Chance within-asset AUC = 0.5",
        fontsize=6.3,
        color=INK,
        ha="left",
        va="top",
        bbox={"facecolor": "white", "edgecolor": "none", "pad": 0.5, "alpha": 0.88},
    )
    bottom.set_ylabel("AUC / AUC loss", fontsize=axis_font)
    bottom.set_xlabel("Planted dynamic signal", fontsize=axis_font)
    bottom.set_title("B. Ordered signal recovery", fontsize=title_font, loc="left", weight="bold")
    bottom.legend(frameon=False, fontsize=legend_font, loc="upper left")

    for ax in axes:
        ax.grid(alpha=0.2, linewidth=0.5)
        ax.tick_params(labelsize=7.4 if full_budget else 7.0)
        ax.spines[["top", "right"]].set_visible(False)
    fig.subplots_adjust(left=0.18, right=0.97, bottom=0.09, top=0.96, hspace=0.52)
    _save_publication_figure(fig, path, "Controlled simulation")
    plt.close(fig)


def _make_appendix_cross_model_figure(path: Path, edition: str) -> None:
    full_budget = edition in {"v3", "v4"}
    results = load_appendix_cross_model_results(
        IDENTITY_DECOMPOSITION_TABLE,
        CROSS_MODEL_RESULTS_TABLE,
    )
    labels = [result.label for result in results]
    pooled = [result.pooled_auc for result in results]
    within = [result.within_asset_auc for result in results]

    fig_size = (3.38, 3.75) if full_budget else (3.35, 3.60)
    fig, ax = plt.subplots(figsize=fig_size, dpi=320)
    y = list(range(len(results)))
    offset = 0.17
    pooled_bars = ax.barh(
        [value + offset for value in y],
        pooled,
        height=0.29,
        color=TEAL,
        label="Pooled",
    )
    within_bars = ax.barh(
        [value - offset for value in y],
        within,
        height=0.29,
        color=GOLD,
        label="Within-asset",
    )
    ax.axvline(CHANCE_AUC, color=INK, linewidth=0.8, linestyle="--", zorder=0)
    ax.text(CHANCE_AUC + 0.003, 5.49, "chance", fontsize=6.2, color=INK, va="bottom")
    ax.set_xlim(0.44, 0.875)
    ax.set_ylim(5.62, -0.55)
    ax.set_xticks([0.5, 0.6, 0.7, 0.8])
    ax.set_xlabel("ROC-AUC", fontsize=8.0 if full_budget else 7.6)
    ax.set_yticks(y, labels, fontsize=7.3 if full_budget else 7.0)
    ax.tick_params(axis="x", labelsize=7.2 if full_budget else 6.9)
    ax.grid(axis="x", alpha=0.2, linewidth=0.5)
    ax.spines[["top", "right", "left"]].set_visible(False)
    fig.suptitle(
        "Cross-model pooled versus\nwithin-asset discrimination",
        fontsize=8.4 if full_budget else 8.0,
        weight="bold",
        color=INK,
        y=0.985,
        linespacing=1.12,
    )
    fig.legend(
        (pooled_bars, within_bars),
        ("Pooled", "Within-asset"),
        frameon=False,
        fontsize=7.2 if full_budget else 6.9,
        loc="upper center",
        bbox_to_anchor=(0.60, 0.925),
        ncol=2,
        columnspacing=1.5,
        handlelength=1.6,
    )
    ax.get_yticklabels()[0].set_fontweight("bold")
    for row, (pooled_value, within_value) in enumerate(zip(pooled, within, strict=True)):
        value_font = 6.6 if full_budget else 6.3
        ax.text(
            pooled_value + 0.006,
            row + offset,
            f"{pooled_value:.3f}",
            va="center",
            fontsize=value_font,
            color=INK,
        )
        ax.text(
            max(within_value + 0.006, 0.507),
            row - offset,
            f"{within_value:.3f}",
            va="center",
            fontsize=value_font,
            color=INK,
            bbox={"facecolor": "white", "edgecolor": "none", "pad": 0.2, "alpha": 0.88},
        )
    fig.subplots_adjust(left=0.36, right=0.97, bottom=0.11, top=0.81)
    _save_publication_figure(
        fig,
        path,
        "Cross-model pooled versus within-asset discrimination",
    )
    plt.close(fig)


def _make_appendix_identity_order_figure(path: Path, edition: str) -> None:
    full_budget = edition in {"v3", "v4"}
    data = load_identity_order_results(
        IDENTITY_DECOMPOSITION_TABLE,
        IDENTITY_SWAP_TABLE,
        TEMPORAL_ORDER_TABLE,
    )
    labels = [result.label for result in data.interventions]
    changes = [result.auc_change for result in data.interventions]
    colours = [RED if result.category == "identity" else TEAL for result in data.interventions]

    fig_size = (3.38, 3.55) if full_budget else (3.35, 3.40)
    fig, ax = plt.subplots(figsize=fig_size, dpi=320)
    y = list(range(len(data.interventions)))
    ax.barh(y, changes, height=0.56, color=colours)
    ax.axvline(0.0, color=INK, linewidth=0.9, zorder=0)
    ax.set_xlim(-0.12, 0.018)
    ax.set_ylim(4.65, -0.55)
    ax.set_xticks([-0.10, -0.05, 0.0])
    ax.set_xlabel("Change in pooled ROC-AUC", fontsize=8.0 if full_budget else 7.6)
    ax.set_yticks(y, labels, fontsize=7.4 if full_budget else 7.0)
    ax.tick_params(axis="x", labelsize=7.2 if full_budget else 6.9)
    ax.grid(axis="x", alpha=0.2, linewidth=0.5)
    ax.spines[["top", "right", "left"]].set_visible(False)
    fig.suptitle(
        "Identity interventions changed ranking more than\norder perturbations",
        fontsize=8.5 if full_budget else 8.1,
        weight="bold",
        color=INK,
        y=0.985,
        linespacing=1.15,
    )
    fig.text(
        0.5,
        0.865,
        f"Original Transformer ROC-AUC = {data.baseline_auc:.6f}",
        ha="center",
        fontsize=6.8 if full_budget else 6.5,
        color=INK,
    )
    fig.legend(
        handles=(
            Patch(facecolor=RED, label="Identity intervention"),
            Patch(facecolor=TEAL, label="Order perturbation"),
        ),
        frameon=False,
        fontsize=6.7 if full_budget else 6.4,
        loc="upper center",
        bbox_to_anchor=(0.5, 0.825),
        ncol=2,
        columnspacing=1.2,
        handlelength=1.4,
    )
    for row, change in enumerate(changes):
        category = data.interventions[row].category
        if category == "identity":
            x, horizontal_alignment, text_colour = change + 0.004, "left", "white"
        else:
            x, horizontal_alignment, text_colour = 0.015, "right", INK
        ax.text(
            x,
            row,
            f"{change:+.4f}",
            ha=horizontal_alignment,
            va="center",
            fontsize=6.7 if full_budget else 6.4,
            color=text_colour,
        )
    fig.subplots_adjust(left=0.42, right=0.96, bottom=0.13, top=0.72)
    _save_publication_figure(
        fig,
        path,
        "Identity interventions changed ranking more than order perturbations",
    )
    plt.close(fig)


def _make_appendix_family_prevalence_figure(path: Path, edition: str) -> None:
    full_budget = edition in {"v3", "v4"}
    data = load_family_prevalence(WINDOW_ENDPOINT_PREVALENCE_TABLE)
    x = np.arange(len(data), dtype=float)
    width = 0.23
    split_styles = (
        ("train", "Train", -width, 1.0, "///", "white"),
        ("validation", "Validation", 0.0, 0.58, "", None),
        ("test", "Test", width, 1.0, "", None),
    )

    fig_size = (3.38, 3.55) if full_budget else (3.35, 3.40)
    fig, ax = plt.subplots(figsize=fig_size, dpi=320)
    for split, _label, offset, alpha, hatch, fixed_face in split_styles:
        for index, row in enumerate(data):
            colour = FAMILY_COLOURS[row.family]
            facecolour = fixed_face or colour
            ax.bar(
                x[index] + offset,
                getattr(row, split),
                width=width,
                facecolor=facecolour,
                edgecolor=colour,
                linewidth=0.8,
                alpha=alpha,
                hatch=hatch,
                zorder=2,
            )

    legend_colour = MID
    fig.legend(
        handles=(
            Patch(
                facecolor="white",
                edgecolor=legend_colour,
                hatch="///",
                label="Train",
            ),
            Patch(
                facecolor=legend_colour,
                edgecolor=legend_colour,
                alpha=0.58,
                label="Validation",
            ),
            Patch(facecolor=legend_colour, edgecolor=legend_colour, label="Test"),
        ),
        frameon=False,
        fontsize=6.8 if full_budget else 6.5,
        loc="upper center",
        bbox_to_anchor=(0.5, 0.845),
        ncol=3,
        columnspacing=1.0,
        handlelength=1.3,
    )
    fig.suptitle(
        "Adverse-event target prevalence\nby family and split",
        fontsize=8.4 if full_budget else 8.0,
        weight="bold",
        color=INK,
        y=0.985,
        linespacing=1.12,
    )
    ax.set_xticks(
        x,
        [FAMILY_LABELS[row.family] for row in data],
        rotation=22,
        ha="right",
        fontsize=6.4 if full_budget else 6.1,
    )
    ax.set_ylabel("Positive-label prevalence", fontsize=7.8 if full_budget else 7.5)
    ax.yaxis.set_major_formatter(PercentFormatter(1.0, decimals=0))
    ax.tick_params(axis="y", labelsize=7.0 if full_budget else 6.7)
    ax.set_ylim(0.0, 0.56)
    ax.grid(axis="y", alpha=0.2, linewidth=0.5)
    ax.spines[["top", "right"]].set_visible(False)

    validation_lookup = {row.family: row.validation for row in data}
    bonds_index = FAMILY_ORDER.index("Bonds")
    bonds_value = validation_lookup["Bonds"]
    ax.annotate(
        f"{bonds_value:.2%}",
        xy=(x[bonds_index], bonds_value),
        xytext=(x[bonds_index] - 0.08, 0.082),
        ha="center",
        va="bottom",
        fontsize=6.2,
        color=INK,
        arrowprops={"arrowstyle": "-", "color": MID, "linewidth": 0.6},
    )
    crypto_index = FAMILY_ORDER.index("Crypto")
    crypto_value = validation_lookup["Crypto"]
    ax.text(
        x[crypto_index],
        crypto_value + 0.014,
        f"{crypto_value:.2%}",
        ha="center",
        va="bottom",
        fontsize=6.2,
        color=INK,
    )
    fig.subplots_adjust(left=0.17, right=0.98, bottom=0.23, top=0.73)
    _save_publication_figure(
        fig,
        path,
        "Adverse-event target prevalence by family and split",
    )
    plt.close(fig)


def _make_appendix_asset_prevalence_figure(path: Path, edition: str) -> None:
    full_budget = edition in {"v3", "v4"}
    data = load_asset_prevalence(WINDOW_ENDPOINT_PREVALENCE_TABLE)
    if data.excluded != MODEL_WINDOW_EXCLUSIONS:
        rendered = "; ".join(f"{item.ticker}: {item.reason}" for item in data.excluded)
        raise ValueError(f"Figure A4 has unexpected configured-universe exclusions: {rendered}")

    fig_size = (3.38, 3.55) if full_budget else (3.35, 3.45)
    fig, ax = plt.subplots(figsize=fig_size, dpi=320)
    for family in FAMILY_ORDER:
        points = [point for point in data.points if point.family == family]
        ax.scatter(
            [point.train for point in points],
            [point.test for point in points],
            s=25,
            alpha=0.78,
            color=FAMILY_COLOURS[family],
            edgecolor="white",
            linewidth=0.35,
            label=FAMILY_LABELS[family].replace("\n", " "),
            zorder=3,
        )

    equality_x = [point[0] for point in PREVALENCE_EQUALITY_LINE]
    equality_y = [point[1] for point in PREVALENCE_EQUALITY_LINE]
    ax.plot(equality_x, equality_y, color=INK, linewidth=0.8, linestyle="--", zorder=1)
    ax.text(0.535, 0.555, "y = x", fontsize=6.2, color=INK, rotation=45)
    ax.set_xlim(0.0, 0.68)
    ax.set_ylim(0.0, 0.68)
    ax.set_aspect("equal", adjustable="box")
    ax.set_xticks([0.0, 0.2, 0.4, 0.6])
    ax.set_yticks([0.0, 0.2, 0.4, 0.6])
    ax.xaxis.set_major_formatter(PercentFormatter(1.0, decimals=0))
    ax.yaxis.set_major_formatter(PercentFormatter(1.0, decimals=0))
    ax.set_xlabel("Training target prevalence", fontsize=7.8 if full_budget else 7.5)
    ax.set_ylabel("Test target prevalence", fontsize=7.8 if full_budget else 7.5)
    ax.tick_params(labelsize=7.0 if full_budget else 6.7)
    ax.grid(alpha=0.2, linewidth=0.5)
    ax.spines[["top", "right"]].set_visible(False)
    fig.suptitle(
        "Training versus test target prevalence by asset",
        fontsize=8.4 if full_budget else 8.0,
        weight="bold",
        color=INK,
        y=0.985,
    )
    legend_handles = [
        Line2D(
            [0],
            [0],
            marker="o",
            linestyle="none",
            markerfacecolor=FAMILY_COLOURS[family],
            markeredgecolor="white",
            markersize=5.0,
            label=FAMILY_LABELS[family].replace("\n", " "),
        )
        for family in FAMILY_ORDER
    ]
    fig.legend(
        handles=legend_handles,
        frameon=False,
        fontsize=5.8 if full_budget else 5.6,
        loc="upper center",
        bbox_to_anchor=(0.5, 0.905),
        ncol=3,
        columnspacing=0.8,
        handletextpad=0.3,
    )
    fig.subplots_adjust(left=0.18, right=0.97, bottom=0.14, top=0.76)
    _save_publication_figure(
        fig,
        path,
        "Training versus test target prevalence by asset",
    )
    plt.close(fig)


def _save_publication_figure(fig, path: Path, title: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(
        path,
        dpi=320,
        facecolor="white",
        metadata={"Title": title},
    )
    fig.savefig(
        path.with_suffix(".svg"),
        format="svg",
        facecolor="white",
        metadata={"Title": title, "Creator": "Matplotlib", "Date": None},
    )


def generate_figures(artifacts: Path, edition: str) -> dict[str, Path]:
    artifacts.mkdir(parents=True, exist_ok=True)
    paths = {
        "methodology": artifacts / "figure_1_methodology.png",
        "core_results": artifacts / "figure_2_core_results.png",
        "simulation": artifacts / "figure_3_simulation.png",
        "appendix_cross_model": artifacts / "figure_A1_cross_model.png",
        "appendix_identity_order": artifacts / "figure_A2_identity_order.png",
        "appendix_family_prevalence": artifacts / "figure_A3_family_prevalence.png",
        "appendix_asset_prevalence": artifacts / "figure_A4_asset_prevalence.png",
    }
    _make_methodology_figure(paths["methodology"], edition)
    _make_core_results_figure(paths["core_results"], edition)
    _make_simulation_figure(paths["simulation"], edition)
    _make_appendix_cross_model_figure(paths["appendix_cross_model"], edition)
    _make_appendix_identity_order_figure(paths["appendix_identity_order"], edition)
    _make_appendix_family_prevalence_figure(paths["appendix_family_prevalence"], edition)
    _make_appendix_asset_prevalence_figure(paths["appendix_asset_prevalence"], edition)
    return paths


def _split_front_matter(lines: list[str]) -> tuple[str, str, str, str, list[str]]:
    if not lines or not lines[0].startswith("# "):
        raise ValueError("Markdown source must begin with an H1 title")
    title = lines[0][2:].strip()
    if title != TITLE:
        raise ValueError(f"Unexpected fixed title: {title!r}")

    author = next((line.split(":", 1)[1].strip() for line in lines if line.startswith("Author:")), "")
    programme = next((line.split(":", 1)[1].strip() for line in lines if line.startswith("Programme:")), "")
    try:
        abstract_index = lines.index("## Abstract")
        first_body_index = next(index for index, line in enumerate(lines) if re.match(r"##\s+1\s", line))
    except (ValueError, StopIteration) as exc:
        raise ValueError("Markdown source needs Abstract and numbered body sections") from exc

    abstract_lines = [line for line in lines[abstract_index + 1 : first_body_index] if line.strip()]
    keywords = next((line for line in abstract_lines if line.startswith("Keywords:")), "")
    abstract = " ".join(line for line in abstract_lines if not line.startswith("Keywords:"))
    return title, author, programme, abstract, [keywords, *lines[first_body_index:]]


def _add_title_and_abstract(doc: Document, title: str, author: str, programme: str, abstract: str, keywords: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    first_line, second_line = title.split(" Forecasting:", 1)
    run = paragraph.add_run(first_line)
    run.add_break(WD_BREAK.LINE)
    run.add_text("Forecasting:" + second_line)
    run.font.name = "Arial"
    run.font.size = Pt(15.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("17252A")

    byline = doc.add_paragraph()
    byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    byline.paragraph_format.space_after = Pt(3)
    byline.paragraph_format.keep_with_next = True
    run = byline.add_run(f"{author} | {programme}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(2)
    heading.paragraph_format.space_after = Pt(1.5)
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run("ABSTRACT")
    run.font.name = "Arial"
    run.font.size = Pt(9.3)
    run.font.bold = True

    abstract_paragraph = doc.add_paragraph()
    abstract_paragraph.paragraph_format.left_indent = Inches(0.22)
    abstract_paragraph.paragraph_format.right_indent = Inches(0.22)
    abstract_paragraph.paragraph_format.space_after = Pt(2)
    abstract_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _add_inline_markdown(abstract_paragraph, abstract, size=8.7)
    _set_keep_together(abstract_paragraph)

    keyword_paragraph = doc.add_paragraph()
    keyword_paragraph.paragraph_format.left_indent = Inches(0.22)
    keyword_paragraph.paragraph_format.right_indent = Inches(0.22)
    keyword_paragraph.paragraph_format.space_after = Pt(3)
    label, values = keywords.split(":", 1)
    run = keyword_paragraph.add_run(label + ":")
    run.bold = True
    run.font.size = Pt(8.2)
    run = keyword_paragraph.add_run(values)
    run.font.size = Pt(8.2)


def _add_figure(doc: Document, key: str, figures: dict[str, Path], edition: str) -> None:
    if key not in figures:
        raise KeyError(f"Unknown figure marker: {key}")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_together = True
    if edition in {"v2", "v3", "v4"}:
        paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(figures[key]), width=Inches(3.31 if edition in {"v2", "v3", "v4"} else 3.25))
    caption = doc.add_paragraph(style="Caption Compact")
    if edition in {"v3", "v4"}:
        caption_text = FIGURE_CAPTIONS_V3[key]
    elif edition == "v2":
        caption_text = FIGURE_CAPTIONS_V2.get(key, FIGURE_CAPTIONS[key])
    else:
        caption_text = FIGURE_CAPTIONS[key]
    _add_inline_markdown(caption, caption_text)


def _add_markdown_table(doc: Document, rows: list[list[str]], table_number: int, edition: str) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.keep_with_next = True
    title.paragraph_format.space_before = Pt(2.5)
    title.paragraph_format.space_after = Pt(1.2)
    captions = TABLE_CAPTIONS_V3 if edition in {"v3", "v4"} else TABLE_CAPTIONS
    run = title.add_run(captions[table_number])
    run.bold = True
    run.font.size = Pt(7.8)

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_width = table._tbl.tblPr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        table._tbl.tblPr.append(tbl_width)
    tbl_width.set(qn("w:w"), "4680")
    tbl_width.set(qn("w:type"), "dxa")
    _set_table_borders(table)
    if edition in {"v3", "v4"} and table_number == 4:
        column_widths = [Inches(0.90), Inches(1.08), Inches(1.26)]
    elif table_number == 1:
        column_widths = [Inches(0.72), Inches(1.15), Inches(1.37)]
    elif len(rows[0]) == 2:
        column_widths = [Inches(0.90), Inches(2.34)]
    elif len(rows[0]) == 3:
        column_widths = [Inches(1.48), Inches(0.88), Inches(0.88)]
    else:
        raise ValueError(f"Unsupported table width: {len(rows[0])} columns")
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Pt(10)
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        if row_index == 0:
            _set_repeat_table_header(row)
        for col_index, value in enumerate(values):
            cell = row.cells[col_index]
            cell.width = column_widths[col_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            if row_index == 0:
                _set_cell_shading(cell, "DCE7E9")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            prose_table = edition in {"v3", "v4"} and table_number in {1, 4}
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_index == 0 or prose_table else WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(7.5 if edition in {"v3", "v4"} and table_number == 4 else 7.8)
            run.bold = row_index == 0


def _flush_paragraph(doc: Document, buffer: list[str], *, references: bool, edition: str) -> None:
    if not buffer:
        return
    text = " ".join(line.strip() for line in buffer)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if references else WD_ALIGN_PARAGRAPH.JUSTIFY
    if references:
        paragraph.paragraph_format.left_indent = Inches(0.13)
        paragraph.paragraph_format.first_line_indent = Inches(-0.13)
        paragraph.paragraph_format.space_after = Pt(0.5)
        _add_inline_markdown(paragraph, text, size=7.8)
    else:
        _add_inline_markdown(paragraph, text)
    if references or edition == "v1" or text.startswith("The aim is to determine"):
        _set_keep_together(paragraph)
    buffer.clear()


def _add_body(doc: Document, lines: list[str], figures: dict[str, Path], edition: str) -> None:
    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    _configure_section(body_section, columns=2)
    footer = doc.sections[0].footer
    if not footer.paragraphs:
        footer.add_paragraph()
    _add_page_number(footer.paragraphs[0])

    paragraph_buffer: list[str] = []
    table_buffer: list[list[str]] = []
    table_number = 0
    in_table = False
    in_references = False

    for raw_line in lines:
        line = raw_line.rstrip()
        if not line.strip():
            if in_table:
                if len(table_buffer) < 2:
                    raise ValueError("Malformed Markdown table")
                table_number += 1
                _add_markdown_table(doc, [table_buffer[0], *table_buffer[2:]], table_number, edition)
                table_buffer.clear()
                in_table = False
            _flush_paragraph(doc, paragraph_buffer, references=in_references, edition=edition)
            continue

        if line.startswith("|"):
            _flush_paragraph(doc, paragraph_buffer, references=in_references, edition=edition)
            values = [cell.strip() for cell in line.strip().strip("|").split("|")]
            table_buffer.append(values)
            in_table = True
            continue

        if in_table:
            table_number += 1
            _add_markdown_table(doc, [table_buffer[0], *table_buffer[2:]], table_number, edition)
            table_buffer.clear()
            in_table = False

        figure_match = re.fullmatch(r"\[\[FIGURE:([a-z_]+)\]\]", line.strip())
        if figure_match:
            _flush_paragraph(doc, paragraph_buffer, references=in_references, edition=edition)
            _add_figure(doc, figure_match.group(1), figures, edition)
            continue

        if line.startswith("## "):
            _flush_paragraph(doc, paragraph_buffer, references=in_references, edition=edition)
            heading = line[3:].strip()
            doc.add_paragraph(heading, style="Heading 1")
            in_references = heading == "References" or heading.endswith(" References")
            continue

        if line.startswith("### "):
            _flush_paragraph(doc, paragraph_buffer, references=in_references, edition=edition)
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            continue

        if line.startswith("> "):
            _flush_paragraph(doc, paragraph_buffer, references=in_references, edition=edition)
            paragraph = doc.add_paragraph(style="Equation Compact")
            source_equation = line[2:].strip()
            equation_render = EQUATION_RENDER_V2 if edition in {"v2", "v3", "v4"} else EQUATION_RENDER
            paragraph.add_run(equation_render.get(source_equation, source_equation))
            continue

        if line.startswith("- "):
            _flush_paragraph(doc, paragraph_buffer, references=in_references, edition=edition)
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.16)
            paragraph.paragraph_format.first_line_indent = Inches(-0.10)
            paragraph.paragraph_format.space_after = Pt(1)
            _add_inline_markdown(paragraph, line[2:].strip())
            continue

        paragraph_buffer.append(line)

    if in_table:
        table_number += 1
        _add_markdown_table(doc, [table_buffer[0], *table_buffer[2:]], table_number, edition)
    _flush_paragraph(doc, paragraph_buffer, references=in_references, edition=edition)


def build(source: Path, output: Path, artifacts: Path, edition: str) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    title, author, programme, abstract, body = _split_front_matter(lines)
    keywords = body.pop(0)
    figures = generate_figures(artifacts, edition)

    doc = Document()
    _configure_document(doc)
    _add_title_and_abstract(doc, title, author, programme, abstract, keywords)
    _add_body(doc, body, figures, edition)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Built {output}")
    print(f"Generated {len(figures)} publication figures under {artifacts}")


def main() -> None:
    args = parse_args()
    source = args.source.resolve()
    if args.edition:
        edition = args.edition
    elif source.stem.endswith("_v4"):
        edition = "v4"
    elif source.stem.endswith("_v3"):
        edition = "v3"
    elif source.stem.endswith("_v2"):
        edition = "v2"
    else:
        edition = "v1"
    build(source, args.output.resolve(), args.artifacts.resolve(), edition)


if __name__ == "__main__":
    main()
